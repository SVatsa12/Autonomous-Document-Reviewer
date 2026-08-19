"""Multi-format document processor supporting PDF, DOCX, and images (PNG, JPEG, TIFF)."""

import io
import os
import tempfile
from pathlib import Path
from typing import Union, List, Optional

# PDF processing
try:
    from pypdf import PdfReader
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

# DOCX processing
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# Image processing (OCR)
try:
    import pytesseract
    from PIL import Image
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

from functions import clean_text


class DocumentProcessor:
    """Domain-agnostic document processor for multiple formats."""
    
    SUPPORTED_EXTENSIONS = {
        '.pdf': 'PDF document (digital or scanned)',
        '.docx': 'Microsoft Word document',
        '.png': 'Image (PNG)',
        '.jpg': 'Image (JPEG)',
        '.jpeg': 'Image (JPEG)',
        '.tiff': 'Image (TIFF)',
        '.tif': 'Image (TIFF)',
    }
    
    def __init__(self, tesseract_path: Optional[str] = None):
        """
        Initialize document processor.
        
        Args:
            tesseract_path: Path to tesseract executable (for OCR)
        """
        if tesseract_path and HAS_OCR:
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
    
    @staticmethod
    def get_supported_formats() -> List[str]:
        """Return list of supported file extensions."""
        return list(DocumentProcessor.SUPPORTED_EXTENSIONS.keys())
    
    @staticmethod
    def is_supported(file_path: Union[str, Path]) -> bool:
        """Check if file format is supported."""
        return Path(file_path).suffix.lower() in DocumentProcessor.SUPPORTED_EXTENSIONS
    
    def process(self, file_path: Union[str, Path]) -> str:
        """
        Process document and extract text.
        
        Args:
            file_path: Path to document file
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If file format not supported
            ImportError: If required dependencies missing
        """
        file_path = Path(file_path)
        suffix = file_path.suffix.lower()
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if suffix == '.pdf':
            return self._process_pdf(file_path)
        elif suffix == '.docx':
            return self._process_docx(file_path)
        elif suffix in ['.png', '.jpg', '.jpeg', '.tiff', '.tif']:
            return self._process_image(file_path)
        else:
            raise ValueError(
                f"Unsupported file format: {suffix}\n"
                f"Supported: {', '.join(DocumentProcessor.SUPPORTED_EXTENSIONS.keys())}"
            )
    
    def _process_pdf(self, file_path: Path) -> str:
        """Extract text from PDF (digital or scanned)."""
        if not HAS_PDF:
            raise ImportError("pypdf not installed. Run: pip install pypdf")
        
        text_parts = []
        reader = PdfReader(str(file_path))
        
        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text and page_text.strip():
                text_parts.append(page_text)
            else:
                # Page might be scanned/image-based - fallback to OCR if available
                if HAS_OCR:
                    try:
                        # Extract images from PDF page
                        from pdf2image import convert_from_path
                        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
                            tmp.write(file_path.read_bytes())
                            tmp_path = tmp.name
                        
                        try:
                            images = convert_from_path(tmp_path, dpi=300)
                            for img in images:
                                ocr_text = pytesseract.image_to_string(img, lang='eng')
                                if ocr_text.strip():
                                    text_parts.append(ocr_text)
                        finally:
                            os.unlink(tmp_path)
                    except Exception as e:
                        print(f"Warning: OCR failed for page {page_num}: {e}")
                else:
                    print(f"Warning: Page {page_num} has no extractable text. Install pytesseract+pypdf2image for OCR.")
        
        full_text = "\n\n".join(text_parts)
        return clean_text(full_text)
    
    def _process_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        if not HAS_DOCX:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        doc = Document(str(file_path))
        text_parts = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        full_text = "\n\n".join(text_parts)
        return clean_text(full_text)
    
    def _process_image(self, file_path: Path) -> str:
        """Extract text from image using OCR."""
        if not HAS_OCR:
            raise ImportError(
                "pytesseract not installed. Run: pip install pytesseract pillow\n"
                "Also install Tesseract OCR engine: https://github.com/tesseract-ocr/tesseract"
            )
        
        try:
            image = Image.open(str(file_path))
            # Convert to RGB if needed (for PNG with transparency)
            if image.mode in ('RGBA', 'LA', 'P'):
                background = Image.new('RGB', image.size, (255, 255, 255))
                if image.mode == 'P':
                    image = image.convert('RGBA')
                background.paste(image, mask=image.split()[-1] if image.mode == 'RGBA' else None)
                image = background
            elif image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Perform OCR
            text = pytesseract.image_to_string(image, lang='eng')
            return clean_text(text)
        except Exception as e:
            raise ValueError(f"Failed to process image {file_path}: {e}")
    
    def process_batch(self, file_paths: List[Union[str, Path]]) -> List[str]:
        """Process multiple documents."""
        results = []
        for fp in file_paths:
            results.append(self.process(fp))
        return results


def detect_domain(file_path: Union[str, Path]) -> str:
    """
    Auto-detect document domain based on filename and content patterns.
    Returns: 'legal', 'banking', 'hr', 'generic', or 'unknown'
    """
    path = Path(file_path)
    filename = path.name.lower()
    
    # Banking/Financial keywords
    banking_keywords = ['loan', 'mortgage', 'credit', 'application', 'kyc', 'aml', 
                        'account', 'statement', 'financial', 'bank', 'tax']
    
    # Legal keywords
    legal_keywords = ['agreement', 'contract', 'lease', 'rental', ' will', 'deed',
                      'legal', 'law', 'court', 'plaintiff', 'defendant']
    
    # HR keywords
    hr_keywords = ['employee', 'employment', 'hr', 'policy', 'handbook', 'offer letter',
                   'confidentiality', 'nda', 'termination', 'resignation']
    
    # Check filename
    for keyword in banking_keywords:
        if keyword in filename:
            return 'banking'
    for keyword in legal_keywords:
        if keyword in filename:
            return 'legal'
    for keyword in hr_keywords:
        if keyword in filename:
            return 'hr'
    
    return 'generic'


def get_domain_config(domain: str) -> dict:
    """
    Get domain-specific configuration including rule sets.
    
    Returns:
        dict with keys: rules, metadata, etc.
    """
    configs = {
        'legal': {
            'default_rules': [
                'separate_numbering',
                'no_merge',
                'full_meaning',
                'formatting',
                'identify_parties',
                'extract_obligations',
            ],
            'common_clause_types': [
                'preamble', 'definitions', 'payment', 'termination', 
                'liability', 'governing_law', 'dispute_resolution'
            ],
        },
        'banking': {
            'default_rules': [
                'extract_financial_amounts',
                'identify_parties',
                'flag_high_risk_terms',
                'check_completeness',
                'extract_dates',
            ],
            'common_clause_types': [
                'loan_amount', 'interest_rate', 'repayment_schedule',
                'collateral', 'default', 'prepayment', 'covenants'
            ],
        },
        'hr': {
            'default_rules': [
                'identify_parties',
                'extract_dates',
                'extract_salary_amounts',
                'check_mandatory_clauses',
                'flag_non_compliant_terms',
            ],
            'common_clause_types': [
                'position', 'salary', 'benefits', 'leave', 'confidentiality',
                'non_compete', 'termination', 'notice_period'
            ],
        },
        'generic': {
            'default_rules': [
                'separate_numbering',
                'full_meaning',
                'check_completeness',
            ],
            'common_clause_types': ['general'],
        },
    }
    
    return configs.get(domain, configs['generic'])


# Convenience function
def extract_text(file_path: Union[str, Path]) -> str:
    """
    Extract text from any supported document format.
    
    Example:
        text = extract_text("contract.pdf")
        text = extract_text("loan_app.jpg")
        text = extract_text("policy.docx")
    """
    processor = DocumentProcessor()
    return processor.process(file_path)


if __name__ == "__main__":
    # Quick test
    import sys
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        print(f"Processing: {file_path}")
        print(f"Detected domain: {detect_domain(file_path)}")
        print("-" * 60)
        text = extract_text(file_path)
        print(f"Extracted {len(text)} characters")
        print("\nFirst 500 chars:")
        print(text[:500])
    else:
        print("Usage: python document_processor.py <file_path>")
